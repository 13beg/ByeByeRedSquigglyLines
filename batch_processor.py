import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import sys
from lingua import Language, LanguageDetectorBuilder

# Supported 5 languages
languages = [
    Language.TURKISH,
    Language.ENGLISH,
    Language.FRENCH,
    Language.GERMAN,
    Language.RUSSIAN
]

# Lingua Dedektörünü Oluştur
detector = LanguageDetectorBuilder.from_languages(*languages).build()

# Lingua to MS Word OpenXML languages codes
LANG_MAP = {
    Language.TURKISH: 'tr-TR',
    Language.ENGLISH: 'en-US',
    Language.FRENCH: 'fr-FR',
    Language.GERMAN: 'de-DE',
    Language.RUSSIAN: 'ru-RU'
}

def set_run_language(run, lang_code):
    """Word OpenXML seviyesinde kelime/run dilini ayarlar."""
    rPr = run._r.get_or_add_rPr()
    lang = OxmlElement('w:lang')
    lang.set(qn('w:val'), lang_code)
    lang.set(qn('w:bidi'), lang_code) # Rusça vb. için uyumluluk
    rPr.append(lang)

def split_into_sentences(text):
    """Seperates texts into sentences"""
    sentences = re.split(r'(?<=[.!?]) +', text)
    return [s for s in sentences if s.strip()]

def get_sentence_dominant_language(sentence):
    """
    Cümlenin hakim dilini ve güven skorunu (0.0 - 1.0) hesaplama.
    """
    confidence_values = detector.compute_language_confidence_values(sentence)
    if not confidence_values:
        return Language.TURKISH, 0.0
    
    top_result = confidence_values[0] # En yüksek olasılıklı dil
    return top_result.language, top_result.value

def detect_word_language(word, sentence_default_lang):
    """
    Kelime bazli dil tespiti yapar.
    3 harften kisa veya rakam içeren şeylerde cümlenin varsayilan dilini korur.
    """
    clean_word = re.sub(r'[^\w\s]', '', word).strip()
    if len(clean_word) < 4 or clean_word.isdigit():
        return sentence_default_lang

    confidence_values = detector.compute_language_confidence_values(clean_word)
    if confidence_values and confidence_values[0].value > 0.60:
        return confidence_values[0].language
    
    return sentence_default_lang

def process_file(file_path):
    doc = docx.Document(file_path)

    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        
        sentences = split_into_sentences(paragraph.text)
        paragraph.text = "" # Paragrafı temizle
        
        for sentence in sentences:
            words = sentence.split(" ")
            
            # Cümle Dili ve Güven Skoru Tespiti
            main_lang, confidence = get_sentence_dominant_language(sentence)
            
            # Threshold = %70 (0.70)
            # Eğer cümlenin ana dili %70'ten yüksek bir güvenle tespit edildiyse:
            is_high_confidence = confidence >= 0.70
            
            sentence_runs = []
            
            for word in words:
                if not word:
                    continue
                
                if is_high_confidence:
                    # Yüksek güvenli cümlede kelimeleri kontrol et: 
                    # Araya başka dilden belirgin bir kelime girmiş mi?
                    word_lang = detect_word_language(word, main_lang)
                else:
                    # Düşük güvenli (karma) cümlede her kelimeyi bağımsız analiz et
                    word_lang = detect_word_language(word, main_lang)
                
                xml_lang_code = LANG_MAP.get(word_lang, 'tr-TR')
                sentence_runs.append((word, xml_lang_code))

            # 2. Yan Yana Aynı Dilde Olan Kelimeleri Birleştir (XML Performansı İçin)
            merged_runs = []
            if sentence_runs:
                curr_words, curr_lang = [sentence_runs[0][0]], sentence_runs[0][1]
                for word, lang_code in sentence_runs[1:]:
                    if lang_code == curr_lang:
                        curr_words.append(word)
                    else:
                        merged_runs.append((" ".join(curr_words), curr_lang))
                        curr_words, curr_lang = [word], lang_code
                merged_runs.append((" ".join(curr_words), curr_lang))

            # 3. Word Belgesine Run Olarak Ekle
            for text_chunk, lang_code in merged_runs:
                run = paragraph.add_run(text_chunk + " ")
                set_run_language(run, lang_code)

    output_path = "fixed_" + file_path
    doc.save(output_path)
    print(f"Analysis is completed in 5 languages: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        process_file(sys.argv[1])
    else:
        print("Kullanim: python cli_fixer.py dosya_adi.docx")
